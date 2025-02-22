import os
import tempfile
from docx import Document
import streamlit as st
from docx2pdf import convert
import platform

# Define template paths
TEMPLATE_DIR = "templates"
template_paths = {
    "AI Automation": "Ai_automation.docx",
    "AI Automation without LPW": "AI Automations Proposal wiithout lpw.docx",
    "Digital Marketing": "DM Proposal.docx",
    "Business Automations": "Business Automations Proposal.docx",
    "IT Consultation": "Contract Agreement.docx"
}

def replace_text_preserve_formatting(doc, replacements):
    """Replace text while preserving formatting"""
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, str(value))

def generate_proposal(proposal_type, client_name, replacements):
    """Generate proposal document with given replacements"""
    try:
        template_path = os.path.join(TEMPLATE_DIR, template_paths[proposal_type])
        doc = Document(template_path)
        replace_text_preserve_formatting(doc, replacements)
        
        temp_dir = tempfile.mkdtemp()
        output_docx = os.path.join(temp_dir, f"{proposal_type}_{client_name}.docx")
        output_pdf = os.path.join(temp_dir, f"{proposal_type}_{client_name}.pdf")
        
        # Save DOCX
        doc.save(output_docx)
        
        # Always read DOCX data
        with open(output_docx, "rb") as docx_file:
            docx_data = docx_file.read()
        
        # Try PDF conversion
        pdf_data = None
        try:
            if platform.system() == "Windows":
                import pythoncom
                pythoncom.CoInitialize()
                convert(output_docx, output_pdf)
                with open(output_pdf, "rb") as pdf_file:
                    pdf_data = pdf_file.read()
        except Exception as e:
            st.warning(f"PDF conversion failed: {str(e)}")
        
        return {
            "docx": (docx_data, f"{proposal_type}_{client_name}.docx", 
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            "pdf": (pdf_data, f"{proposal_type}_{client_name}.pdf", "application/pdf") if pdf_data else None
        }
            
    except Exception as e:
        st.error(f"Error generating proposal: {str(e)}")
        return None
    finally:
        try:
            if os.path.exists(temp_dir):
                import shutil
                shutil.rmtree(temp_dir)
        except:
            pass
